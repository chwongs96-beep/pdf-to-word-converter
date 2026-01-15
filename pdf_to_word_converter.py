"""
PDF 到 Word 转换器 - 支持文字识别、转换、编辑和关键词搜索
支持普通PDF和扫描版PDF的文字识别（OCR）
"""
import os
from pathlib import Path
from typing import List, Tuple
from pdf2docx import Converter
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import PyPDF2

# OCR相关导入（可选） - 支持文字识别和提取"""
    
    def __init__(self):
        self.converted_file = None
        self.ocr_available = OCR_AVAILABLE
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class PDFToWordConverter:
    """PDF 转 Word 转换器类"""
    
    def __init__(self):
        self.converted_file = None
    heck_pdf_has_text(self, pdf_path: str) -> bool:
        """
        检查PDF是否包含可提取的文字（非扫描版）
        
        Args:
            pdf_path: PDF 文件路径
            
        Returns:
            True 如果包含文字，False 如果是纯图片/扫描版
        """
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                # 检查前几页
                pages_to_check = min(3, len(pdf_reader.pages))
                
                for i in range(pages_to_check):
                    page = pdf_reader.pages[i]
                    text = page.extract_text().strip()
                    if text and len(text) > 50:  # 如果有足够的文字内容
                        return True
                return False
        except:
            return False
    
    def convert_pdf_to_word(self, pdf_path: str, word_path: str = None, use_ocr: bool = None) -> str:
        """
        将 PDF 文件转换为 Word 文档（智能识别PDF类型并提取文字）
        
        Args:
            pdf_path: PDF 文件路径
            word_path: 输出的 Word 文件路径（可选）
            use_ocr: 是否使用OCR（None=自动检测，True=强制OCR，False=不使用OCR）
            
        Returns:
            生成的 Word 文件路径
        """
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")
        
        # 如果未指定输出路径，使用相同名称但扩展名为 .docx
        if word_path is None:
            word_path = str(Path(pdf_path).with_suffix('.docx'))
        
        # 自动检测是否需要OCR
        if use_ocr is None:
            has_text = self.check_pdf_has_text(pdf_path)
            use_ocr = not has_text
            if use_ocr:
                print(f"检测到扫描版PDF，将使用OCR识别文字...")
            else:
                print(f"检测到文本型PDF，直接提取文字...")
        
        # 如果需要OCR但不可用
        if use_ocr and not self.ocr_available:
            print("警告: OCR功能未安装，将尝试直接转换...")
            print("提示: 安装 pytesseract 和 pdf2image 以支持扫描版PDF")
            use_ocr = False
        
        try:
            if use_ocr:
                # 使用OCR方式
                return self._convert_with_ocr(pdf_path, word_path)
            else:
                # 使用标准方式
                print(f"正在转换 {pdf_path} 到 {word_path}...")
                print("正在提取PDF文字内容...")
                cv = Converter(pdf_path)
                cv.convert(word_path)
                cv.close()
                print(f"✓ 文字提取和转换完成！文件保存在: {word_path}")
                self.converted_file = word_path
                return word_path
        except Exception as e:
            raise Exception(f"转换失败: {str(e)}")
    
    def _convert_with_ocr(self, pdf_path: str, word_path: str) -> str:
        """
        使用OCR方式转换扫描版PDF（识别图片中的文字）
        
        Args:
            pdf_path: PDF 文件路径
            word_path: 输出的 Word 文件路径
            
        Returns:
            生成的 Word 文件路径
        """
        print(f"正在使用OCR识别 {pdf_path} 中的文字...")
        
        # 将PDF转换为图片
        print("步骤 1/3: 将PDF页面转换为图片...")
        images = convert_from_path(pdf_path, dpi=300)
        
        # 创建新的Word文档
        doc = Document()
        
        # 对每一页进行OCR识别
        print(f"步骤 2/3: 识别文字（共 {len(images)} 页）...")
        for i, image in enumerate(images, 1):
            print(f"  正在识别第 {i}/{len(images)} 页...")
            
            # 使用Tesseract进行OCR
            # 支持中文和英文
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            
            # 添加到Word文档
            if i > 1:
                doc.add_page_break()
            
            doc.add_paragraph(text)
        
        # 保存Word文档
        print(f"步骤 3/3: 保存文档...")
        doc.save(word_path)
        print(f"✓ OCR识别和转换完成！文件保存在: {word_path}")
        
        self.converted_file = word_path
        return word_path
            raise Exception(f"转换失败: {str(e)}")
    
    def search_keyword(self, word_path: str, keyword: str) -> List[Tuple[int, str]]:
        """
        在 Word 文档中搜索关键词
        
        Args:
            word_path: Word 文件路径
            keyword: 要搜索的关键词
            
        Returns:
            包含关键词的段落列表 [(段落索引, 段落内容), ...]
        """
        if not os.path.exists(word_path):
            raise FileNotFoundError(f"Word 文件不存在: {word_path}")
        
        doc = Document(word_path)
        results = []
        
        for idx, para in enumerate(doc.paragraphs):
            if keyword.lower() in para.text.lower():
                results.append((idx, para.text))
        
        return results
    
    def highlight_keyword(self, word_path: str, keyword: str, output_path: str = None):
        """
        在 Word 文档中高亮显示关键词
        
        Args:
            word_path: Word 文件路径
            keyword: 要高亮的关键词
            output_path: 输出文件路径（可选）
        """
        if not os.path.exists(word_path):
            raise FileNotFoundError(f"Word 文件不存在: {word_path}")
        
        doc = Document(word_path)
        
        for para in doc.paragraphs:
            if keyword.lower() in para.text.lower():
                for run in para.runs:
                    if keyword.lower() in run.text.lower():
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # 保存文件
        if output_path is None:
            output_path = word_path.replace('.docx', '_highlighted.docx')
        
        doc.save(output_path)
        print(f"已高亮关键词 '{keyword}'，保存在: {output_path}")
        return output_path
    
    def replace_text(self, word_path: str, old_text: str, new_text: str, 
                     output_path: str = None) -> str:
        """
        替换 Word 文档中的文本
        
        Args:
            word_path: Word 文件路径
            old_text: 要替换的文本
            new_text: 新文本
            output_path: 输出文件路径（可选）
            
        Returns:
            输出文件路径
        """
        if not os.path.exists(word_path):
            raise FileNotFoundError(f"Word 文件不存在: {word_path}")
        
        doc = Document(word_path)
        replacement_count = 0
        
        # 替换段落中的文本
        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replacement_count += 1
        
        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    replacement_count += 1
        
        # 保存文件
        if output_path is None:
            output_path = word_path.replace('.docx', '_edited.docx')
        
        doc.save(output_path)
        print(f"已替换 {replacement_count} 处文本，保存在: {output_path}")
        return output_path
    
    def add_text_to_document(self, word_path: str, text: str, 
                            output_path: str = None) -> str:
        """
        在 Word 文档末尾添加新文本
        
        Args:
            word_path: Word 文件路径
            text: 要添加的文本
            output_path: 输出文件路径（可选）
            
        Returns:
            输出文件路径
        """
        if not os.path.exists(word_path):
            raise FileNotFoundError(f"Word 文件不存在: {word_path}")
        
        doc = Document(word_path)
        doc.add_paragraph(text)
        
        if output_path is None:
            output_path = word_path
        
        doc.save(output_path)
        print(f"已添加文本，保存在: {output_path}")
        return output_path
    
    def get_document_info(self, word_path: str) -> dict:
        """
        获取 Word 文档信息
        
        Args:
            word_path: Word 文件路径
            
        Returns:
            文档信息字典
        """
        if not os.path.exists(word_path):
            raise FileNotFoundError(f"Word 文件不存在: {word_path}")
        
        doc = Document(word_path)
        
        info = {
            '段落数': len(doc.paragraphs),
            '表格数': len(doc.tables),
            '总字符数': sum(len(para.text) for para in doc.paragraphs),
        }
        
        return info


def main():
    """主函数 - 演示如何使用转换器"""
    converter = PDFToWordConverter()
    
    print("=" * 50)
    print("PDF 到 Word 转换器")
    print("=" * 50)
    print()
    
    while True:
        print("\n请选择操作：")
        print("1. PDF 转 Word")
        print("2. 搜索关键词")
        print("3. 高亮关键词")
        print("4. 替换文本")
        print("5. 添加文本")
        print("6. 查看文档信息")
        print("0. 退出")
        
        choice = input("\n请输入选项 (0-6): ").strip()
        
        if choice == '0':
            print("谢谢使用！")
            break
        
        elif choice == '1':
            pdf_path = input("请输入 PDF 文件路径: ").strip()
            word_path = input("请输入输出 Word 文件路径 (直接回车使用默认): ").strip()
            word_path = word_path if word_path else None
            
            try:
                result = converter.convert_pdf_to_word(pdf_path, word_path)
                print(f"\n✓ 转换成功！")
            except Exception as e:
                print(f"\n✗ 错误: {e}")
        
        elif choice == '2':
            word_path = input("请输入 Word 文件路径: ").strip()
            keyword = input("请输入要搜索的关键词: ").strip()
            
            try:
                results = converter.search_keyword(word_path, keyword)
                if results:
                    print(f"\n找到 {len(results)} 处包含 '{keyword}' 的内容：")
                    for idx, (para_idx, text) in enumerate(results[:10], 1):
                        print(f"\n{idx}. 段落 {para_idx}:")
                        print(f"   {text[:100]}...")
                    if len(results) > 10:
                        print(f"\n... 还有 {len(results) - 10} 处结果")
                else:
                    print(f"\n未找到关键词 '{keyword}'")
            except Exception as e:
                print(f"\n✗ 错误: {e}")
        
        elif choice == '3':
            word_path = input("请输入 Word 文件路径: ").strip()
            keyword = input("请输入要高亮的关键词: ").strip()
            output_path = input("请输入输出文件路径 (直接回车使用默认): ").strip()
            output_path = output_path if output_path else None
            
            try:
                result = converter.highlight_keyword(word_path, keyword, output_path)
                print(f"\n✓ 高亮完成！")
            except Exception as e:
                print(f"\n✗ 错误: {e}")
        
        elif choice == '4':
            word_path = input("请输入 Word 文件路径: ").strip()
            old_text = input("请输入要替换的文本: ").strip()
            new_text = input("请输入新文本: ").strip()
            output_path = input("请输入输出文件路径 (直接回车使用默认): ").strip()
            output_path = output_path if output_path else None
            
            try:
                result = converter.replace_text(word_path, old_text, new_text, output_path)
                print(f"\n✓ 替换完成！")
            except Exception as e:
                print(f"\n✗ 错误: {e}")
        
        elif choice == '5':
            word_path = input("请输入 Word 文件路径: ").strip()
            text = input("请输入要添加的文本: ").strip()
            
            try:
                result = converter.add_text_to_document(word_path, text)
                print(f"\n✓ 添加完成！")
            except Exception as e:
                print(f"\n✗ 错误: {e}")
        
        elif choice == '6':
            word_path = input("请输入 Word 文件路径: ").strip()
            
            try:
                info = converter.get_document_info(word_path)
                print("\n文档信息：")
                for key, value in info.items():
                    print(f"  {key}: {value}")
            except Exception as e:
                print(f"\n✗ 错误: {e}")
        
        else:
            print("\n✗ 无效选项，请重新选择")


if __name__ == "__main__":
    main()
